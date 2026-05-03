"""DOCX loader -- python-docx alapon, natív szöveg + táblázat-kinyerés.

A DOCX mindig digitális (NEM szkennelt), tehát egyszerűbb mint a PDF —
nincs OCR/vision fallback. A táblázatokat Markdown formátumba alakítjuk a
`tables_markdown` mezőhöz.
"""

from __future__ import annotations

from io import BytesIO

from graph.states.pipeline_state import IngestedDocument, PageContent


def load_docx(file_name: str, file_bytes: bytes) -> IngestedDocument:
    """Egy DOCX betöltése IngestedDocument-té (mindig digitális, egy oldal)."""
    import docx

    try:
        doc = docx.Document(BytesIO(file_bytes))
    except Exception as e:
        raise RuntimeError(f"Nem sikerult megnyitni a DOCX-et: {file_name}: {e}") from e

    # Bekezdések szövege
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]

    # Táblázatok Markdown-ba
    table_blocks: list[str] = []
    table_count = 0
    for tbl_idx, tbl in enumerate(doc.tables, start=1):
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in tbl.rows]
        rows = [r for r in rows if any(c for c in r)]
        if not rows:
            continue
        n_cols = max(len(r) for r in rows)
        if n_cols == 0:
            continue
        # Header
        header = list(rows[0]) + [""] * (n_cols - len(rows[0]))
        sep = ["---"] * n_cols
        body = []
        for r in rows[1:]:
            padded = list(r) + [""] * (n_cols - len(r))
            body.append("| " + " | ".join(c[:30] for c in padded[:n_cols]) + " |")
        md = (
            "| " + " | ".join(c[:30] for c in header[:n_cols]) + " |\n"
            "| " + " | ".join(sep) + " |\n"
            + "\n".join(body)
        )
        table_blocks.append(f"### Táblázat #{tbl_idx}\n\n{md}\n")
        table_count += 1

    full_text = "\n\n".join(paragraphs)
    tables_markdown = "\n".join(table_blocks)

    return IngestedDocument(
        file_name=file_name,
        file_type="docx",
        pages=[PageContent(page_number=1, text=full_text, is_scanned=False)],
        full_text=full_text,
        tables_markdown=tables_markdown,
        table_count=table_count,
        is_scanned=False,
    )
