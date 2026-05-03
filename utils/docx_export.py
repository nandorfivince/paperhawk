"""DOCX report generation via python-docx.

10-section structure:
  1. Title + date
  2. Metadata (provider, model, processing time)
  3. Performance metrics (manual estimate vs speedup)
  4. Executive summary (LLM-generated when available)
  5. Documents table (file_name, doc_type, evidence_score)
  6. Cross-document checks (three-way matching)
  7. Risks color-coded (red / orange / blue)
  8. Package-level analysis (when state["package_insights"])
  9. DD analysis (when state["dd_report"])
  10. Footer (applied standards list)

python-docx is blocking; the caller (export_docx_node) wraps it in
``asyncio.to_thread``.
"""

from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor

from graph.states.pipeline_state import (
    DDPortfolioReport,
    PackageInsights,
    PipelineState,
)


# Severity color codes
_COLOR_HIGH = RGBColor(0xCC, 0x00, 0x00)
_COLOR_MEDIUM = RGBColor(0xCC, 0x88, 0x00)
_COLOR_LOW = RGBColor(0x00, 0x33, 0x99)
_COLOR_INFO = RGBColor(0x66, 0x66, 0x66)


def _color_for(severity: str) -> RGBColor:
    return {
        "high": _COLOR_HIGH,
        "medium": _COLOR_MEDIUM,
        "low": _COLOR_LOW,
        "info": _COLOR_INFO,
    }.get(severity.lower(), _COLOR_INFO)


def build_docx_sync(state: PipelineState) -> bytes:
    """Sync DOCX builder. The caller invokes via ``asyncio.to_thread()``."""
    doc = Document()

    # 1. Title
    title = doc.add_heading("Agentic Document Intelligence — Audit Report", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

    # 2. Metadata
    meta = doc.add_paragraph()
    meta.add_run("Generated at: ").bold = True
    meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))

    docs_count = len(state.get("documents") or [])
    meta = doc.add_paragraph()
    meta.add_run("Documents processed: ").bold = True
    meta.add_run(str(docs_count))

    # 3. Performance metrics
    report = state.get("report") or {}
    perf = report.get("performance") or {}
    if perf:
        doc.add_heading("Performance metrics", level=2)
        p = doc.add_paragraph()
        p.add_run(
            f"Processing time: {perf.get('processing_seconds', 0):.2f} sec | "
            f"Manual estimate: {perf.get('manual_estimate_minutes', 0)} min | "
            f"Speedup: {perf.get('speedup', 0):.1f}x"
        ).bold = True

    # 4. Executive summary
    if report.get("executive_summary"):
        doc.add_heading("Executive summary", level=2)
        doc.add_paragraph(report["executive_summary"])

    # 5. Documents table
    docs_info = report.get("documents") or []
    if docs_info:
        doc.add_heading("Documents", level=2)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Light Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "File"
        hdr[1].text = "Type"
        hdr[2].text = "Fields"
        hdr[3].text = "Evidence (ISA 500)"
        for d in docs_info:
            row = tbl.add_row().cells
            row[0].text = str(d.get("file", ""))
            row[1].text = str(d.get("type", ""))
            row[2].text = str(d.get("extracted_fields", 0))
            row[3].text = f"{d.get('evidence_score', 0)}/10"

    # 6. Cross-document checks
    comparison = state.get("comparison")
    if comparison:
        doc.add_heading("Cross-document checks", level=2)
        p = doc.add_paragraph()
        p.add_run(
            f"Checks: {comparison.total_checks} -- "
            f"{comparison.ok_count} ok, {comparison.warning_count} warnings, "
            f"{comparison.critical_count} critical, {comparison.missing_count} missing."
        ).italic = True

        # Show only non-ok mismatches
        non_ok = [m for m in comparison.matches if m.get("severity") != "ok"]
        if non_ok:
            for m in non_ok:
                sev = m.get("severity", "warning")
                msg = m.get("message", "")
                prefix = {
                    "critical": "CRITICAL",
                    "warning": "WARNING",
                }.get(sev, sev.upper())
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"{prefix}: {msg}")
                run.font.color.rgb = _COLOR_HIGH if sev == "critical" else _COLOR_MEDIUM

    # 7. Risks color-coded
    risks = state.get("risks") or []
    if risks:
        doc.add_heading("Risks", level=2)
        for severity in ("high", "medium", "low", "info"):
            sev_risks = [r for r in risks if r.severity.lower() == severity]
            if not sev_risks:
                continue
            sub = doc.add_heading(severity.upper(), level=3)
            for run in sub.runs:
                run.font.color.rgb = _color_for(severity)
            for r in sev_risks[:20]:  # max 20 per category
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(r.description)
                run.font.color.rgb = _color_for(severity)
                if r.rationale:
                    p.add_run(f" — {r.rationale}").italic = True

    # 8. Package-level analysis
    pkg: PackageInsights | None = state.get("package_insights")
    if pkg:
        doc.add_heading("Package-level analysis", level=2)
        doc.add_paragraph(pkg.executive_summary or "")
        if pkg.findings:
            for f in pkg.findings:
                if isinstance(f, dict):
                    doc.add_paragraph(
                        f.get("description") or f.get("leiras", ""),
                        style="List Bullet",
                    )

    # 9. DD analysis
    dd: DDPortfolioReport | None = state.get("dd_report")
    if dd:
        doc.add_heading("DD analysis (contract portfolio)", level=2)
        doc.add_paragraph(f"Contract count: {dd.contract_count}")
        if dd.executive_summary:
            doc.add_paragraph(dd.executive_summary)
        if dd.top_red_flags:
            doc.add_heading("Top red flags", level=3)
            for flag in dd.top_red_flags:
                doc.add_paragraph(flag, style="List Bullet")
        if dd.total_monthly_obligations:
            doc.add_heading("Monthly obligations (estimated)", level=3)
            for cur, amt in dd.total_monthly_obligations.items():
                doc.add_paragraph(f"{cur}: {amt:,.0f}")
        if dd.contracts:
            doc.add_heading("Per-contract risk level", level=3)
            tbl = doc.add_table(rows=1, cols=4)
            tbl.style = "Light Grid"
            hdr = tbl.rows[0].cells
            hdr[0].text = "File"
            hdr[1].text = "Type"
            hdr[2].text = "Parties"
            hdr[3].text = "Risk"
            for c in dd.contracts:
                if hasattr(c, "model_dump"):
                    c = c.model_dump()
                row = tbl.add_row().cells
                row[0].text = str(c.get("file_name", ""))
                row[1].text = str(c.get("contract_type", ""))
                row[2].text = ", ".join(c.get("parties") or [])
                level = (c.get("risk_level") or "low").upper()
                run = row[3].paragraphs[0].add_run(level)
                run.bold = True
                run.font.color.rgb = _color_for(level.lower())
                # Red flags (if any)
                red_flags = c.get("red_flags") or []
                if red_flags:
                    p = doc.add_paragraph()
                    p.add_run(f"  Red flags ({c.get('file_name','')}): ").bold = True
                    p.add_run("; ".join(red_flags[:5]))
        if dd.expiring_soon:
            doc.add_heading("Expiring soon (within 12 months)", level=3)
            for fname in dd.expiring_soon:
                doc.add_paragraph(fname, style="List Bullet")

    # 10. Footer — only the actually applied standards
    from domain_checks import get_applied_standards
    standards = get_applied_standards(risks) if risks else []
    doc.add_paragraph()
    foot = doc.add_paragraph()
    if standards:
        foot.add_run(
            f"Applied standards and methods: {' | '.join(standards)}"
        ).font.size = Pt(8)
    else:
        foot.add_run(
            "Generated by: Agentic Document Intelligence Platform (LangGraph)."
        ).font.size = Pt(8)

    # Bytes
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
