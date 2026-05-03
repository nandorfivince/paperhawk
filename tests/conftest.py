"""Pytest fixtures — used across the whole test suite."""

from __future__ import annotations

import pytest


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    """Bytes for a minimal English invoice PDF (PyMuPDF-rendered).

    The full ``test_data/generate_samples.py`` produces much richer files; this
    fixture exists for ingest-level unit tests so they don't depend on the
    full ``test_data/`` regeneration.
    """
    import fitz

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)  # A4
    text = (
        "INVOICE\n\n"
        "Invoice number: 2026/001\n"
        "Issue date: 2026-01-31\n\n"
        "Issuer: AcmeSoft Inc.\n"
        "Tax ID: 12-3456789\n\n"
        "Customer: BudaData LLC\n"
        "Tax ID: 98-7654321\n\n"
        "Line items:\n"
        "Software development services   40 hours   $500.00   $20,000.00\n\n"
        "Total net: $20,000.00\n"
        "Total VAT: $4,000.00 (20%)\n"
        "Total gross: $24,000.00\n"
    )
    page.insert_text((50, 50), text, fontsize=11)
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


@pytest.fixture
def sample_docx_bytes() -> bytes:
    """Bytes for a minimal English contract DOCX."""
    import io

    import docx

    doc = docx.Document()
    doc.add_heading("Non-Disclosure Agreement", level=1)
    doc.add_paragraph(
        "Parties: SmartSensors Inc. (tax id: 13-5792468) and "
        "InfoTech Ltd. (tax id: 86-4201357)"
    )
    doc.add_paragraph("Effective date: 2026-01-15")
    doc.add_paragraph("Expiry date: 2027-01-15")
    doc.add_paragraph(
        "Penalty: A breach of this confidentiality obligation triggers a $50,000 penalty per incident."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def sample_png_bytes() -> bytes:
    """Bytes for a minimal PNG (white background + caption)."""
    import io

    from PIL import Image, ImageDraw

    img = Image.new("RGB", (800, 600), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 50), "Invoice test PNG", fill="black")

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()
