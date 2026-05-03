"""Synthetic sample-file generation for the test data set.

Generates PDF + DOCX + PNG files in business-style ID-based file names
(NOT finding-cheat names). The sample documents are crafted so each known
risk pattern is provoked exactly once — the system must detect it from the
content, not from the filename.

Layout:
  * invoices/           — 3 EN invoices (audit pattern: March 50% pricier)
                          + 1 EN intra-EU + 1 DE Rechnung (multilingual demo)
  * contracts/          — 1 NDA + 1 MSSA + 1 IT framework + 1 DE→HU lease
                          (multilingual demo includes HU/DE elements)
  * multi_doc/          — invoice + delivery_note + purchase_order with
                          quantity discrepancy (40 vs 38)
  * demo_packages/      — pre-built demo bundles for the pitch:
      audit_demo/      — 3 invoices, March 50% pricier
      dd_demo/         — NDA + MSSA (3 red flags) + amendment
      compliance_demo/ — 2 contracts; one missing GDPR Article 28
  * adversarial/        — 4 deliberately broken docs (math, incomplete, bilingual, dates)
  * financial_reports/  — 1 EN income statement (US-GAAP) + 1 EN cash flow (IFRS)

Run:  python test_data/generate_samples.py
"""

from __future__ import annotations

import random
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from PIL import Image

random.seed(2026)

THIS_DIR = Path(__file__).parent

# Hungarian tax-id CDV (mod-10) — for the optional HU multilingual fixture
_HU_TAX_WEIGHTS = [9, 7, 3, 1, 9, 7, 3]


def _compute_cdv(first7: str) -> int:
    checksum = sum(int(d) * w for d, w in zip(first7[:7], _HU_TAX_WEIGHTS, strict=False))
    return (10 - (checksum % 10)) % 10


def _make_hu_tax(first8: str, region: str = "42") -> str:
    cdv = _compute_cdv(first8[:7])
    return f"{first8[:7]}{cdv}-2-{region}"


def _money(amount: float | int, currency: str = "USD") -> str:
    """US-style money: '1,234,567.00 USD'."""
    if currency == "HUF":
        return f"{amount:,.0f}".replace(",", " ") + " Ft"
    if currency == "EUR":
        return f"{amount:,.2f} EUR"
    return f"${amount:,.2f}"


# Companies (EN-first)
COMPANIES = {
    "AcmeSoft": {"tax_id": "12-3456789", "address": "100 Market St, New York, NY 10001, USA"},
    "DataPharm": {"tax_id": "98-7654321", "address": "200 Mission St, San Francisco, CA 94105, USA"},
    "PestTrade": {"tax_id": "24-6802468", "address": "500 King St W, Toronto, ON M5V 1L9, Canada"},
    "BorgenLab": {"tax_id": "13-5792468", "address": "75 Park Lane, London W1K 1RA, UK"},
    "NorthTech": {"tax_id": "86-4201357", "address": "120 Adelaide St E, Toronto, ON M5C 1K9, Canada"},
    "BuilderInc": {"tax_id": "11-2233445", "address": "1500 Industrial Blvd, Chicago, IL 60616, USA"},
    "ConstructLLC": {"tax_id": "55-6677889", "address": "850 Riverside Dr, Houston, TX 77002, USA"},
    "TechSupply": {"tax_id": "21-4365879", "address": "300 Beach Ave, Los Angeles, CA 90001, USA"},
    "AcmeBuy": {"tax_id": "65-7483920", "address": "60 Wall St, New York, NY 10005, USA"},
    "GlobalCorp": {"tax_id": "33-4455667", "address": "100 Federal St, Boston, MA 02110, USA"},
    "MediCare": {"tax_id": "77-8899001", "address": "200 Pearl St, Hartford, CT 06103, USA"},
    "DataVendor": {"tax_id": "99-0011223", "address": "1 Market St, Dallas, TX 75202, USA"},
    "CleanLaw": {"tax_id": "44-5566778", "address": "500 Boylston St, Boston, MA 02116, USA"},
    "MullerBauer": {"tax_id": "31-4159265", "address": "1000 Wilshire Blvd, Los Angeles, CA 90017, USA"},
}

# Multilingual fallback fixtures (HU, DE — for multilingual demo proof)
HU_COMPANIES = {
    "BudaSoft": {"tax_id": _make_hu_tax("12345678", "42"), "address": "1137 Budapest, Szent István krt. 12., Hungary"},
    "DataPharmHU": {"tax_id": _make_hu_tax("98765432", "41"), "address": "1095 Budapest, Lechner Ödön fasor 9., Hungary"},
    "EpitoKft": {"tax_id": _make_hu_tax("11223344", "13"), "address": "1221 Budapest, Építő utca 1., Hungary"},
    "VarEpito": {"tax_id": _make_hu_tax("55667788", "42"), "address": "1221 Budapest, Nagytétényi út 190., Hungary"},
}

EU_COMPANIES = {
    "NLLogistics": {
        "name": "Netherlands Logistics B.V.",
        "vat_id": "NL854321987B01",
        "address": "Prins Hendrikkade 21, 1012 TL Amsterdam, Netherlands",
    },
    "BavarianKraftwerk": {
        "name": "Bavarian Kraftwerk GmbH",
        "vat_id": "DE123456789",
        "address": "Maximilianstraße 12, 80539 München, Germany",
    },
    "AlpenTech": {
        "name": "AlpenTech AG",
        "vat_id": "DE987654321",
        "address": "Königstraße 30, 70173 Stuttgart, Germany",
    },
}


def _render_html_pdf(out_path: Path, html: str) -> None:
    """HTML → A4 PDF via PyMuPDF."""
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    rect = fitz.Rect(40, 40, 555, 802)
    full_html = f"""<!doctype html><html><head><meta charset="utf-8"><style>
        body {{ font-family: sans-serif; font-size: 10pt; color: #000; }}
        h1 {{ font-size: 18pt; margin: 0 0 8pt 0; }}
        h2 {{ font-size: 12pt; margin: 12pt 0 4pt 0; }}
        p  {{ margin: 4pt 0; }}
        table {{ width: 100%; border-collapse: collapse; margin: 6pt 0; }}
        th, td {{ border: 1px solid #444; padding: 3pt 5pt; text-align: left; }}
        th {{ background: #ddd; }}
        .right {{ text-align: right; }}
        .total {{ font-weight: bold; background: #eef; }}
    </style></head><body>{html}</body></html>"""
    page.insert_htmlbox(rect, full_html)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path), garbage=4, deflate=True)
    doc.close()


def _render_docx(out_path: Path, sections: list[tuple[str, str]]) -> None:
    doc = Document()
    for h, b in sections:
        if h:
            doc.add_heading(h, level=1)
        if b:
            for para in b.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def _render_png_from_pdf(pdf_path: Path, png_path: Path, dpi: int = 200) -> None:
    doc = fitz.open(str(pdf_path))
    page = doc[0]
    pix = page.get_pixmap(dpi=dpi)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    png_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(png_path, "PNG")
    doc.close()


# ---------------------------------------------------------------------------
# Invoice template (EN)
# ---------------------------------------------------------------------------


def _invoice_html(
    inv_no: str, issue: str, fulfillment: str, due: str,
    issuer: str, issuer_tax: str, issuer_addr: str,
    customer: str, customer_tax: str, customer_addr: str,
    line_items: list[dict],
    currency: str = "USD",
) -> str:
    net = sum(t["net"] for t in line_items)
    vat = sum(t["net"] * t["vat_pct"] / 100 for t in line_items)
    gross = net + vat
    rows = "\n".join(
        f"<tr><td>{t['name']}</td>"
        f"<td class='right'>{t['quantity']}</td>"
        f"<td class='right'>{_money(t['unit_price'], currency)}</td>"
        f"<td class='right'>{_money(t['net'], currency)}</td>"
        f"<td class='right'>{t['vat_pct']}%</td></tr>"
        for t in line_items
    )
    return f"""
    <h1>INVOICE</h1>
    <p><b>Invoice number:</b> {inv_no}</p>
    <p><b>Issue date:</b> {issue} &nbsp;&nbsp;
       <b>Fulfillment date:</b> {fulfillment} &nbsp;&nbsp;
       <b>Payment due:</b> {due}</p>
    <h2>Issuer</h2>
    <p>{issuer}<br/>Tax ID: {issuer_tax}<br/>Address: {issuer_addr}</p>
    <h2>Customer</h2>
    <p>{customer}<br/>Tax ID: {customer_tax}<br/>Address: {customer_addr}</p>
    <h2>Line items</h2>
    <table>
        <tr><th>Description</th><th>Quantity</th><th>Unit price</th><th>Net</th><th>VAT</th></tr>
        {rows}
    </table>
    <table>
        <tr><td><b>Total net</b></td><td class="right">{_money(net, currency)}</td></tr>
        <tr><td><b>Total VAT</b></td><td class="right">{_money(vat, currency)}</td></tr>
        <tr class="total"><td><b>Total gross</b></td><td class="right">{_money(gross, currency)}</td></tr>
    </table>
    """


def _invoice_docx_sections(
    inv_no: str, dates: dict, parties: dict, line_items: list[dict],
    net: float, vat: float, gross: float, currency: str = "USD",
) -> list[tuple[str, str]]:
    items_text = "\n".join(
        f"{t['name']} -- {t['quantity']} units -- {_money(t['unit_price'], currency)}/unit -- "
        f"net {_money(t['net'], currency)} -- VAT {t['vat_pct']}%"
        for t in line_items
    )
    return [
        ("INVOICE", f"Invoice number: {inv_no}\nIssue date: {dates['issue']}\n"
                    f"Fulfillment date: {dates['fulfillment']}\nPayment due: {dates['due']}"),
        ("Issuer", f"{parties['issuer']}\nTax ID: {parties['issuer_tax']}\nAddress: {parties['issuer_addr']}"),
        ("Customer", f"{parties['customer']}\nTax ID: {parties['customer_tax']}\nAddress: {parties['customer_addr']}"),
        ("Line items", items_text),
        ("Totals",
         f"Total net: {_money(net, currency)}\nTotal VAT: {_money(vat, currency)}\n"
         f"Total gross: {_money(gross, currency)}"),
    ]


# ---------------------------------------------------------------------------
# Generators
# ---------------------------------------------------------------------------


def generate_invoices() -> None:
    """3 EN invoices showing an audit-pattern price increase + multilingual fallback."""
    base_price = 500.00  # USD/hour
    common = {
        "issuer": "AcmeSoft Inc.",
        "issuer_tax": COMPANIES["AcmeSoft"]["tax_id"],
        "issuer_addr": COMPANIES["AcmeSoft"]["address"],
        "customer": "DataPharm LLC",
        "customer_tax": COMPANIES["DataPharm"]["tax_id"],
        "customer_addr": COMPANIES["DataPharm"]["address"],
    }

    invoices = [
        {"no": "2026/001", "issue": "2026-01-31", "fulfillment": "2026-01-30", "due": "2026-02-28",
         "qty": 40, "price": base_price, "out": "as-2026-001.pdf"},
        {"no": "2026/002", "issue": "2026-02-28", "fulfillment": "2026-02-27", "due": "2026-03-30",
         "qty": 42, "price": round(base_price * 1.05, 2), "out": "as-2026-002.pdf"},
        {"no": "2026/003", "issue": "2026-03-31", "fulfillment": "2026-03-29", "due": "2026-04-30",
         "qty": 44, "price": round(base_price * 1.50, 2), "out": "as-2026-003.pdf"},
    ]

    for inv in invoices:
        net = inv["qty"] * inv["price"]
        line_items = [{
            "name": "Software development services",
            "quantity": inv["qty"],
            "unit_price": inv["price"],
            "net": net,
            "vat_pct": 20,
        }]
        vat = net * 0.20
        gross = net + vat

        pdf_path = THIS_DIR / "invoices" / inv["out"]
        html = _invoice_html(
            inv_no=inv["no"], issue=inv["issue"], fulfillment=inv["fulfillment"], due=inv["due"],
            line_items=line_items, **common,
        )
        _render_html_pdf(pdf_path, html)

        docx_path = pdf_path.with_suffix(".docx")
        sections = _invoice_docx_sections(
            inv_no=inv["no"],
            dates={"issue": inv["issue"], "fulfillment": inv["fulfillment"], "due": inv["due"]},
            parties=common,
            line_items=line_items,
            net=net, vat=vat, gross=gross,
        )
        _render_docx(docx_path, sections)
        _render_png_from_pdf(pdf_path, pdf_path.with_suffix(".png"))


def generate_intra_eu_invoice() -> None:
    """EN intra-EU invoice with 0% VAT (reverse charge, Art. 138 EU VAT Directive).

    Tests the false-positive filter: 0% VAT alone is NOT a risk in EU intra-Community context.
    """
    sender = EU_COMPANIES["NLLogistics"]
    buyer_name = "PestTrade Ltd."
    buyer_vat = COMPANIES["PestTrade"]["tax_id"]
    buyer_addr = COMPANIES["PestTrade"]["address"]

    inv_no = "NL-INV-2026-0001"
    issue_date = "2026-02-15"
    due_date = "2026-03-15"
    delivery_date = "2026-02-12"
    net_eur = 6_200
    vat_eur = 0  # intra-EU reverse charge

    html = f"""
    <h1>INVOICE</h1>
    <p><b>Invoice number:</b> {inv_no}</p>
    <p><b>Issue date:</b> {issue_date} &nbsp;&nbsp;
       <b>Delivery date:</b> {delivery_date} &nbsp;&nbsp;
       <b>Payment due:</b> {due_date}</p>
    <h2>Supplier</h2>
    <p>{sender['name']}<br/>VAT ID: {sender['vat_id']}<br/>{sender['address']}</p>
    <h2>Customer</h2>
    <p>{buyer_name}<br/>Tax ID: {buyer_vat}<br/>{buyer_addr}</p>
    <h2>Line items</h2>
    <table>
        <tr><th>Description</th><th>Qty</th><th>Unit price</th><th>Net total</th><th>VAT</th></tr>
        <tr>
            <td>Intra-EU freight forwarding services (Amsterdam-Toronto)</td>
            <td class='right'>1</td>
            <td class='right'>{net_eur:,} EUR</td>
            <td class='right'>{net_eur:,} EUR</td>
            <td class='right'>0%</td>
        </tr>
    </table>
    <table>
        <tr><td><b>Net total</b></td><td class='right'>{net_eur:,} EUR</td></tr>
        <tr><td><b>VAT (0% — Intra-Community supply, reverse charge per Art. 138 EU VAT Directive 2006/112/EC)</b></td><td class='right'>{vat_eur} EUR</td></tr>
        <tr class='total'><td><b>Gross total</b></td><td class='right'>{net_eur:,} EUR</td></tr>
    </table>
    <p><i>Payment terms: 30 days net. Bank: ABN AMRO, IBAN NL12ABNA0123456789.</i></p>
    """
    pdf_path = THIS_DIR / "invoices" / "nl-inv-2026-0001.pdf"
    _render_html_pdf(pdf_path, html)
    _render_docx(pdf_path.with_suffix(".docx"), [
        ("INVOICE",
         f"Invoice number: {inv_no}\nIssue date: {issue_date}\n"
         f"Delivery date: {delivery_date}\nPayment due: {due_date}"),
        ("Supplier", f"{sender['name']}\nVAT ID: {sender['vat_id']}\n{sender['address']}"),
        ("Customer", f"{buyer_name}\nTax ID: {buyer_vat}\n{buyer_addr}"),
        ("Line items",
         f"Intra-EU freight forwarding services -- 1 unit -- {net_eur:,} EUR -- "
         f"VAT 0% (Intra-Community supply, Art. 138 EU VAT Directive)"),
        ("Totals",
         f"Net total: {net_eur:,} EUR\nVAT: 0 EUR (reverse charge)\nGross total: {net_eur:,} EUR"),
    ])
    _render_png_from_pdf(pdf_path, pdf_path.with_suffix(".png"))


def generate_de_rechnung() -> None:
    """DE Rechnung (multilingual demo): 19% MwSt, German language detection."""
    sender = EU_COMPANIES["BavarianKraftwerk"]
    buyer = EU_COMPANIES["AlpenTech"]

    rechnung_no = "BK-R-2026-0001"
    rechnungsdatum = "15.02.2026"
    leistungsdatum = "10.02.2026"
    zahlbar_bis = "17.03.2026"
    netto_eur = 4_800
    mwst_pct = 19
    mwst_eur = round(netto_eur * mwst_pct / 100, 2)
    brutto_eur = netto_eur + mwst_eur

    html = f"""
    <h1>RECHNUNG</h1>
    <p><b>Rechnungsnummer:</b> {rechnung_no}</p>
    <p><b>Rechnungsdatum:</b> {rechnungsdatum} &nbsp;&nbsp;
       <b>Leistungsdatum:</b> {leistungsdatum} &nbsp;&nbsp;
       <b>Zahlbar bis:</b> {zahlbar_bis}</p>
    <h2>Lieferant</h2>
    <p>{sender['name']}<br/>USt-IdNr.: {sender['vat_id']}<br/>{sender['address']}</p>
    <h2>Empfänger</h2>
    <p>{buyer['name']}<br/>USt-IdNr.: {buyer['vat_id']}<br/>{buyer['address']}</p>
    <h2>Leistungen</h2>
    <table>
        <tr><th>Beschreibung</th><th>Menge</th><th>Einzelpreis</th><th>Netto</th><th>MwSt</th></tr>
        <tr>
            <td>Industrieanlagen-Wartung (Q1/2026)</td>
            <td class='right'>1</td>
            <td class='right'>{netto_eur:,} EUR</td>
            <td class='right'>{netto_eur:,} EUR</td>
            <td class='right'>{mwst_pct}%</td>
        </tr>
    </table>
    <table>
        <tr><td><b>Nettobetrag</b></td><td class='right'>{netto_eur:,} EUR</td></tr>
        <tr><td><b>MwSt {mwst_pct}%</b></td><td class='right'>{mwst_eur:,} EUR</td></tr>
        <tr class='total'><td><b>Bruttobetrag</b></td><td class='right'>{brutto_eur:,} EUR</td></tr>
    </table>
    <p><i>Zahlungsbedingungen: 30 Tage netto. Bank: HypoVereinsbank, IBAN DE89370400440532013000.</i></p>
    """
    pdf_path = THIS_DIR / "invoices" / "bk-r-2026-0001.pdf"
    _render_html_pdf(pdf_path, html)
    _render_docx(pdf_path.with_suffix(".docx"), [
        ("RECHNUNG",
         f"Rechnungsnummer: {rechnung_no}\nRechnungsdatum: {rechnungsdatum}\n"
         f"Leistungsdatum: {leistungsdatum}\nZahlbar bis: {zahlbar_bis}"),
        ("Lieferant", f"{sender['name']}\nUSt-IdNr.: {sender['vat_id']}\n{sender['address']}"),
        ("Empfänger", f"{buyer['name']}\nUSt-IdNr.: {buyer['vat_id']}\n{buyer['address']}"),
        ("Leistungen",
         f"Industrieanlagen-Wartung (Q1/2026) -- 1 -- {netto_eur:,} EUR -- MwSt {mwst_pct}%"),
        ("Beträge",
         f"Nettobetrag: {netto_eur:,} EUR\nMwSt {mwst_pct}%: {mwst_eur:,} EUR\n"
         f"Bruttobetrag: {brutto_eur:,} EUR"),
    ])
    _render_png_from_pdf(pdf_path, pdf_path.with_suffix(".png"))


def generate_contracts() -> None:
    """1 NDA + 1 MSSA (clean) + 1 IT framework + 1 DE→HU lease (multilingual demo)."""
    contracts_dir = THIS_DIR / "contracts"

    # 1) NDA — clean (no red flags)
    nda_html = f"""
    <h1>NON-DISCLOSURE AGREEMENT (NDA)</h1>
    <p><b>Parties:</b> BorgenLab Ltd. (tax id: {COMPANIES['BorgenLab']['tax_id']},
       {COMPANIES['BorgenLab']['address']}) and NorthTech Inc. (tax id: {COMPANIES['NorthTech']['tax_id']},
       {COMPANIES['NorthTech']['address']}).</p>
    <p><b>Effective date:</b> 2026-01-15 &nbsp;&nbsp; <b>Expiry date:</b> 2027-01-15</p>
    <h2>1. Scope of confidential information</h2>
    <p>All technical, business, and financial data shared between the parties under this
       agreement, including software specifications, customer lists, and pricing models.</p>
    <h2>2. Confidentiality term</h2>
    <p>The receiving party shall keep the disclosed information confidential for 5 years
       after the expiry of this agreement.</p>
    <h2>3. Penalty</h2>
    <p>Each breach of the confidentiality obligation shall trigger a contractual penalty
       of $50,000 per incident.</p>
    <h2>4. Governing law</h2>
    <p>This agreement shall be governed by the laws of the State of Delaware, USA.</p>
    """
    nda_path = contracts_dir / "bl-nt-nda-2026.pdf"
    _render_html_pdf(nda_path, nda_html)
    _render_docx(nda_path.with_suffix(".docx"), [
        ("NON-DISCLOSURE AGREEMENT",
         f"Parties: BorgenLab Ltd. (tax id: {COMPANIES['BorgenLab']['tax_id']}) and "
         f"NorthTech Inc. (tax id: {COMPANIES['NorthTech']['tax_id']})\n\n"
         "Effective date: 2026-01-15\nExpiry date: 2027-01-15\n\n"
         "Penalty: $50,000 per breach.\n\n"
         "Governing law: State of Delaware, USA."),
    ])
    _render_png_from_pdf(nda_path, nda_path.with_suffix(".png"))

    # 2) MSSA (Master Software Service Agreement) — clean
    mssa_html = f"""
    <h1>MASTER SOFTWARE SERVICE AGREEMENT</h1>
    <p><b>Parties:</b> PestTrade Ltd. (tax id: {COMPANIES['PestTrade']['tax_id']}) as Provider,
       and DataPharm LLC (tax id: {COMPANIES['DataPharm']['tax_id']}) as Client.</p>
    <p><b>Effective date:</b> 2026-02-01 &nbsp;&nbsp; <b>Expiry date:</b> 2027-01-31</p>
    <p><b>Monthly fee:</b> $20,000 + 20% VAT (gross $24,000)</p>
    <h2>1. Scope of services</h2>
    <p>Operation of a cloud-based data analytics platform with 99.5% monthly SLA.</p>
    <h2>2. Change of control</h2>
    <p>If a 50% or greater ownership change occurs at the Provider, the Client shall be
       entitled to terminate this agreement with immediate effect.</p>
    <h2>3. Auto-renewal</h2>
    <p>This agreement automatically renews for an additional one-year term unless either
       party provides written notice of non-renewal at least 60 days before expiry.</p>
    <h2>4. Penalty</h2>
    <p>For each 1% of SLA shortfall, a penalty of $1,000 is due.</p>
    <h2>5. Governing law</h2>
    <p>State of New York, USA.</p>
    """
    mssa_path = contracts_dir / "pt-dp-mssa-2026.pdf"
    _render_html_pdf(mssa_path, mssa_html)
    _render_docx(mssa_path.with_suffix(".docx"), [
        ("MASTER SOFTWARE SERVICE AGREEMENT",
         f"Parties: PestTrade Ltd. (tax id: {COMPANIES['PestTrade']['tax_id']}) and "
         f"DataPharm LLC (tax id: {COMPANIES['DataPharm']['tax_id']})\n\n"
         "Effective date: 2026-02-01\nExpiry date: 2027-01-31\nMonthly fee: $20,000 + 20% VAT\n\n"
         "Change of control: 50% ownership change → immediate termination right.\n\n"
         "Auto-renewal: 1-year term with 60-day notice.\n\n"
         "Penalty: $1,000 per 1% SLA shortfall.\n\n"
         "Governing law: State of New York, USA."),
    ])
    _render_png_from_pdf(mssa_path, mssa_path.with_suffix(".png"))

    # 3) IT framework agreement with 200% SLA penalty (industry-standard, NOT a risk)
    mb = COMPANIES["MullerBauer"]
    dp = COMPANIES["DataPharm"]
    framework_html = f"""
    <h1>IT FRAMEWORK AGREEMENT</h1>
    <p><b>Parties:</b><br/>
       MullerBauer Inc. (tax id: {mb['tax_id']}, registered at: {mb['address']}) as Provider, and<br/>
       DataPharm LLC (tax id: {dp['tax_id']}, registered at: {dp['address']}) as Client.</p>
    <p><b>Effective date:</b> 2026-01-01 &nbsp;&nbsp; <b>Expiry date:</b> 2028-12-31</p>

    <h2>1. Scope</h2>
    <p>The Provider delivers IT support under this framework: infrastructure monitoring,
       incident handling (24/7), patch management, security updates. Custom work via
       individual statements of work, billed at $280/hour + VAT.</p>

    <h2>2. SLA (Service Level Agreement)</h2>
    <p><b>Availability:</b> 99.5% monthly.<br/>
       <b>Reaction time (P1 incident):</b> 30 minutes.<br/>
       <b>Resolution time (P1):</b> 4 business hours.<br/>
       <b>SLA breach penalty:</b> the Provider owes a penalty of
       <b>200% of the affected monthly retainer</b> (industry-standard sanction in
       the IT/SaaS sector).</p>

    <h2>3. Termination</h2>
    <p>Either party may terminate this agreement with 60 days' written notice.
       Material breach permits immediate termination upon written notice with a
       14-day cure period.</p>

    <h2>4. Confidentiality</h2>
    <p>The parties shall keep all information shared under this agreement confidential
       for 5 years. A breach triggers a $100,000 penalty.</p>

    <h2>5. Data protection (GDPR Article 28)</h2>
    <p>The Provider acts as data processor on the Client's documented instructions.
       Processing covers system logs collected as part of infrastructure monitoring.
       No data is transferred to third countries. The Client has audit rights once
       per year. The Provider holds an ISO 27001 certification (since 2018).</p>

    <h2>6. Governing law and jurisdiction</h2>
    <p>State of California, USA. The parties submit to the exclusive jurisdiction of
       the federal courts of the Northern District of California.</p>

    <h2>7. Payment terms</h2>
    <p>Monthly retainer: $18,000 + 20% VAT, payment terms: net 30 days.</p>
    """
    framework_path = contracts_dir / "mbk-it-fa-2026.pdf"
    _render_html_pdf(framework_path, framework_html)
    _render_docx(framework_path.with_suffix(".docx"), [
        ("IT FRAMEWORK AGREEMENT",
         f"Parties: MullerBauer Inc. (tax id: {mb['tax_id']}) as Provider, "
         f"DataPharm LLC (tax id: {dp['tax_id']}) as Client.\n\n"
         "Effective date: 2026-01-01\nExpiry date: 2028-12-31"),
        ("Scope",
         "IT support framework: monitoring (24/7), incident handling, patch management, "
         "security updates. Custom work at $280/hour + VAT."),
        ("SLA",
         "Availability: 99.5% monthly.\nReaction time (P1): 30 minutes.\n"
         "Resolution time (P1): 4 business hours.\n"
         "SLA breach: 200% of the monthly retainer (industry-standard sanction)."),
        ("Termination",
         "60 days' written notice. Material breach: immediate, with 14-day cure period."),
        ("Confidentiality",
         "5-year confidentiality term. Breach: $100,000 penalty."),
        ("GDPR Article 28",
         "Provider as data processor. Subject: system logs from infrastructure monitoring. "
         "No third-country transfers. Annual audit rights. ISO 27001 certified."),
        ("Governing law",
         "State of California, USA. Federal courts of the Northern District of California."),
        ("Payment", "Monthly retainer: $18,000 + 20% VAT. Net 30 days."),
    ])
    _render_png_from_pdf(framework_path, framework_path.with_suffix(".png"))

    # 4) DE→HU lease (multilingual demo): EUR/month, mixed German + English context
    de = EU_COMPANIES["BavarianKraftwerk"]  # leasing-style entity
    lease_html = f"""
    <h1>EQUIPMENT LEASE / LEASINGVERTRAG</h1>
    <p><b>Lessor / Leasinggeber:</b><br/>
       Deutsche Fleet Leasing GmbH (USt-IdNr.: DE556677889,
       Theatinerstraße 8, 80333 München, Germany).</p>
    <p><b>Lessee / Leasingnehmer:</b><br/>
       Budapest Logistics Kft. (tax id: {_make_hu_tax("47852136", "42")},
       1097 Budapest, Könyves Kálmán krt. 12-14., Hungary).</p>
    <p><b>Effective / Vertragsbeginn:</b> 2026-03-01 &nbsp;&nbsp;
       <b>Expiry / Vertragsende:</b> 2029-02-28 (36 months)</p>

    <h2>1. Leased asset</h2>
    <p>15 × Mercedes-Benz Sprinter 316 CDI (model year 2025) on operating lease.
       VIN list and technical specs in a separate annex. Use: domestic and intra-EU goods transport.</p>

    <h2>2. Lease fees</h2>
    <p><b>Down payment:</b> 18,500 EUR.<br/>
       <b>Monthly lease:</b> 1,850 EUR + 27% VAT (HU local VAT, since use is on HU territory).<br/>
       <b>Residual value:</b> 22,000 EUR at end of term.<br/>
       <b>Payment method:</b> SEPA, 5th day of each month.</p>

    <h2>3. Termination</h2>
    <p>Extraordinary termination during the term is permitted only on material breach or
       insolvency of the Lessee. Ordinary termination is permitted from month 24, with
       60 days' notice.</p>

    <h2>4. Penalty</h2>
    <p>Late payment: 5% annual interest. Early termination outside of clause 3:
       60% of remaining lease payments become immediately due.</p>

    <h2>5. Maintenance</h2>
    <p>The Lessor provides full maintenance and casco insurance (Vollkasko mit
       500 EUR Selbstbeteiligung). Fuel and tires at the Lessee's expense.</p>

    <h2>6. GDPR Article 28</h2>
    <p>The fleet management telematics system (GPS, driving data) processes personal data of
       the Lessee's employees. Purpose: vehicle position tracking and maintenance scheduling.
       Retention: 24 months. The Lessee is the controller, the Lessor is the processor under
       a data processing addendum.</p>

    <h2>7. Governing law / Anwendbares Recht</h2>
    <p>German BGB and Hungarian Civil Code apply jointly under conflict-of-laws rules
       (place-of-performance jurisdiction prevails). Disputes resolved by the courts
       of Munich and Budapest jointly.</p>
    """
    lease_path = contracts_dir / "df-lc-2026.pdf"
    _render_html_pdf(lease_path, lease_html)
    _render_docx(lease_path.with_suffix(".docx"), [
        ("EQUIPMENT LEASE / LEASINGVERTRAG",
         "Lessor: Deutsche Fleet Leasing GmbH (USt-IdNr.: DE556677889)\n"
         f"Lessee: Budapest Logistics Kft. (tax id: {_make_hu_tax('47852136', '42')})\n\n"
         "Effective: 2026-03-01 — Expiry: 2029-02-28 (36 months)"),
        ("Asset", "15 × Mercedes-Benz Sprinter 316 CDI (operating lease) for intra-EU goods transport."),
        ("Fees",
         "Down payment: 18,500 EUR\nMonthly lease: 1,850 EUR + 27% VAT\n"
         "Residual: 22,000 EUR\nPayment: SEPA, 5th of month"),
        ("Termination",
         "Extraordinary: material breach / insolvency.\nOrdinary: from month 24, 60 days' notice."),
        ("Penalty",
         "Late: 5% annual interest.\nEarly termination: 60% of remaining lease payments due."),
        ("Maintenance",
         "Lessor: full maintenance + Vollkasko (500 EUR Selbstbeteiligung).\n"
         "Lessee: fuel + tires."),
        ("GDPR Article 28",
         "Telematics GPS system with personal data (drivers).\n"
         "Controller: Lessee. Processor: Lessor.\n"
         "Retention: 24 months."),
        ("Governing law",
         "German BGB + Hungarian Civil Code (conflict-of-laws). "
         "Munich + Budapest courts."),
    ])
    _render_png_from_pdf(lease_path, lease_path.with_suffix(".png"))


def generate_multi_doc_triplet() -> None:
    """Three-way matching: PO (40 units) + delivery_note (38 units, shortage) + invoice (40 units, over-billing)."""
    multi_dir = THIS_DIR / "multi_doc"
    supplier = "BuilderInc Inc."
    supplier_tax = COMPANIES["BuilderInc"]["tax_id"]
    supplier_addr = COMPANIES["BuilderInc"]["address"]
    customer = "ConstructLLC LLC"
    customer_tax = COMPANIES["ConstructLLC"]["tax_id"]
    customer_addr = COMPANIES["ConstructLLC"]["address"]

    # Purchase order: 40 units of HI-100 at $185.00/unit
    po_html = f"""
    <h1>PURCHASE ORDER</h1>
    <p><b>PO number:</b> PO-2026/0412 &nbsp;&nbsp; <b>Date:</b> 2026-04-01</p>
    <p><b>Delivery due date:</b> 2026-04-15</p>
    <h2>Supplier</h2>
    <p>{supplier}<br/>Tax ID: {supplier_tax}<br/>Address: {supplier_addr}</p>
    <h2>Customer</h2>
    <p>{customer}<br/>Tax ID: {customer_tax}<br/>Address: {customer_addr}</p>
    <h2>Line items</h2>
    <table>
        <tr><th>Item code</th><th>Description</th><th>Quantity</th><th>Unit price</th><th>Net</th></tr>
        <tr><td>HI-100</td><td>I-beam 6m</td><td class="right">40</td><td class="right">$185.00</td><td class="right">$7,400.00</td></tr>
    </table>
    """
    po_path = multi_dir / "bi-po-2026-0412.pdf"
    _render_html_pdf(po_path, po_html)
    _render_docx(po_path.with_suffix(".docx"), [
        ("PURCHASE ORDER",
         f"PO number: PO-2026/0412\nDate: 2026-04-01\nDelivery due: 2026-04-15\n\n"
         f"Supplier: {supplier} (tax id: {supplier_tax})\n"
         f"Customer: {customer} (tax id: {customer_tax})\n\n"
         "Line items:\nHI-100 I-beam 6m -- 40 units -- $185.00/unit -- net $7,400.00"),
    ])
    _render_png_from_pdf(po_path, po_path.with_suffix(".png"))

    # Delivery note: 38 units (2 short)
    dn_html = f"""
    <h1>DELIVERY NOTE</h1>
    <p><b>Delivery note number:</b> DN-2026/0415 &nbsp;&nbsp; <b>Date:</b> 2026-04-14</p>
    <p><b>PO reference:</b> PO-2026/0412</p>
    <h2>Supplier</h2><p>{supplier}<br/>Tax ID: {supplier_tax}</p>
    <h2>Customer</h2><p>{customer}<br/>Tax ID: {customer_tax}</p>
    <h2>Line items</h2>
    <table>
        <tr><th>Item code</th><th>Description</th><th>Quantity</th></tr>
        <tr><td>HI-100</td><td>I-beam 6m</td><td class="right">38 units</td></tr>
    </table>
    <p><b>Notes:</b> Due to inventory shortage, 38 units delivered out of the 40 ordered.
       The remaining 2 units will arrive with the next shipment.</p>
    """
    dn_path = multi_dir / "bi-dn-2026-0415.pdf"
    _render_html_pdf(dn_path, dn_html)
    _render_docx(dn_path.with_suffix(".docx"), [
        ("DELIVERY NOTE",
         f"Delivery note number: DN-2026/0415\nDate: 2026-04-14\nPO reference: PO-2026/0412\n\n"
         f"Supplier: {supplier}\nCustomer: {customer}\n\n"
         "Line items:\nHI-100 I-beam 6m -- 38 units (2 units short)"),
    ])
    _render_png_from_pdf(dn_path, dn_path.with_suffix(".png"))

    # Invoice: 40 units (over-billing — should match delivery_note 38 instead)
    net = 40 * 185.00
    vat = net * 0.20
    gross = net + vat
    inv_html = _invoice_html(
        inv_no="2026/BI-0418",
        issue="2026-04-18", fulfillment="2026-04-14", due="2026-05-18",
        issuer=supplier, issuer_tax=supplier_tax, issuer_addr=supplier_addr,
        customer=customer, customer_tax=customer_tax, customer_addr=customer_addr,
        line_items=[{"name": "HI-100 I-beam 6m", "quantity": 40, "unit_price": 185.00,
                     "net": net, "vat_pct": 20}],
    )
    inv_path = multi_dir / "bi-inv-2026-0418.pdf"
    _render_html_pdf(inv_path, inv_html)
    _render_docx(inv_path.with_suffix(".docx"), _invoice_docx_sections(
        inv_no="2026/BI-0418",
        dates={"issue": "2026-04-18", "fulfillment": "2026-04-14", "due": "2026-05-18"},
        parties={"issuer": supplier, "issuer_tax": supplier_tax, "issuer_addr": supplier_addr,
                 "customer": customer, "customer_tax": customer_tax, "customer_addr": customer_addr},
        line_items=[{"name": "HI-100 I-beam 6m", "quantity": 40, "unit_price": 185.00,
                     "net": net, "vat_pct": 20}],
        net=net, vat=vat, gross=gross,
    ))
    _render_png_from_pdf(inv_path, inv_path.with_suffix(".png"))


# ---------------------------------------------------------------------------
# Demo packages
# ---------------------------------------------------------------------------


def generate_audit_demo() -> None:
    """Audit demo: 3 invoices from the same supplier; March is 50% pricier."""
    out_dir = THIS_DIR / "demo_packages" / "audit_demo"
    common = {
        "issuer": "TechSupply Inc.",
        "issuer_tax": COMPANIES["TechSupply"]["tax_id"],
        "issuer_addr": COMPANIES["TechSupply"]["address"],
        "customer": "AcmeBuy Corp.",
        "customer_tax": COMPANIES["AcmeBuy"]["tax_id"],
        "customer_addr": COMPANIES["AcmeBuy"]["address"],
    }

    invoices = [
        {"no": "TS-2026/0101", "issue": "2026-01-31", "fulfillment": "2026-01-30", "due": "2026-02-28",
         "qty": 10, "price": 787.40, "out": "ts-2026-0101.pdf"},
        {"no": "TS-2026/0228", "issue": "2026-02-28", "fulfillment": "2026-02-27", "due": "2026-03-30",
         "qty": 10, "price": 826.77, "out": "ts-2026-0228.pdf"},
        {"no": "TS-2026/0331", "issue": "2026-03-31", "fulfillment": "2026-03-29", "due": "2026-04-30",
         "qty": 10, "price": 1240.16, "out": "ts-2026-0331.pdf"},
    ]

    for inv in invoices:
        net = inv["qty"] * inv["price"]
        line_items = [{
            "name": "Maintenance services (monthly retainer)",
            "quantity": inv["qty"],
            "unit_price": inv["price"],
            "net": net,
            "vat_pct": 20,
        }]
        vat = net * 0.20
        gross = net + vat

        pdf_path = out_dir / inv["out"]
        html = _invoice_html(
            inv_no=inv["no"], issue=inv["issue"], fulfillment=inv["fulfillment"], due=inv["due"],
            line_items=line_items, **common,
        )
        _render_html_pdf(pdf_path, html)

        docx_path = pdf_path.with_suffix(".docx")
        sections = _invoice_docx_sections(
            inv_no=inv["no"],
            dates={"issue": inv["issue"], "fulfillment": inv["fulfillment"], "due": inv["due"]},
            parties=common,
            line_items=line_items,
            net=net, vat=vat, gross=gross,
        )
        _render_docx(docx_path, sections)
        _render_png_from_pdf(pdf_path, pdf_path.with_suffix(".png"))


def generate_dd_demo() -> None:
    """DD demo: NDA + service agreement (3 red flags) + amendment."""
    out_dir = THIS_DIR / "demo_packages" / "dd_demo"

    # 1) NDA — clean
    nda_html = f"""
    <h1>NON-DISCLOSURE AGREEMENT (NDA)</h1>
    <p><b>Parties:</b> GlobalCorp Inc. (tax id: {COMPANIES['GlobalCorp']['tax_id']},
       {COMPANIES['GlobalCorp']['address']}) and DataVendor LLC (tax id: {COMPANIES['DataVendor']['tax_id']},
       {COMPANIES['DataVendor']['address']}).</p>
    <p><b>Effective date:</b> 2026-03-01 &nbsp;&nbsp; <b>Expiry date:</b> 2028-03-01</p>
    <h2>1. Confidential information</h2>
    <p>Technical, business, and financial information mutually shared between the parties.</p>
    <h2>2. Confidentiality term</h2>
    <p>2 years after expiry of this agreement.</p>
    <h2>3. Governing law</h2>
    <p>State of Delaware, USA.</p>
    """
    nda_path = out_dir / "gc-dv-nda-2026-0301.pdf"
    _render_html_pdf(nda_path, nda_html)
    _render_docx(nda_path.with_suffix(".docx"), [
        ("NON-DISCLOSURE AGREEMENT",
         f"Parties: GlobalCorp Inc. (tax id: {COMPANIES['GlobalCorp']['tax_id']}) and "
         f"DataVendor LLC (tax id: {COMPANIES['DataVendor']['tax_id']})\n\n"
         "Effective date: 2026-03-01\nExpiry date: 2028-03-01\n\n"
         "Confidentiality term: 2 years post-expiry.\nGoverning law: State of Delaware, USA."),
    ])
    _render_png_from_pdf(nda_path, nda_path.with_suffix(".png"))

    # 2) MSSA with 3 red flags: change-of-control + non-compete + auto-renewal
    mssa_html = f"""
    <h1>MASTER SOFTWARE SERVICE AGREEMENT</h1>
    <p><b>Parties:</b> DataVendor LLC (tax id: {COMPANIES['DataVendor']['tax_id']}) as Provider,
       and GlobalCorp Inc. (tax id: {COMPANIES['GlobalCorp']['tax_id']}) as Client.</p>
    <p><b>Effective date:</b> 2026-03-15 &nbsp;&nbsp; <b>Expiry date:</b> 2027-03-15</p>
    <p><b>Monthly fee:</b> $35,000 + 20% VAT (gross $42,000)</p>
    <h2>1. Scope</h2>
    <p>Operation of a cloud-based data analytics platform with 99.9% SLA.</p>
    <h2>2. Change of control</h2>
    <p>If a 25% or greater ownership change occurs at the Provider, the Client shall be
       entitled to terminate this agreement with immediate effect, and the Provider
       shall pay the full annual fee ($420,000) as a contractual penalty.</p>
    <h2>3. Non-compete</h2>
    <p>For 5 years after the termination of this agreement the Provider shall not provide
       similar services to the Client's competitors in the United States territory.</p>
    <h2>4. Auto-renewal</h2>
    <p>This agreement automatically renews for an additional 3-year term unless either
       party provides written notice of non-renewal at least 90 days before expiry.</p>
    <h2>5. Penalty</h2>
    <p>For each 1% of SLA shortfall, a $2,000 penalty is due.</p>
    """
    mssa_path = out_dir / "gc-dv-mssa-2026-0315.pdf"
    _render_html_pdf(mssa_path, mssa_html)
    _render_docx(mssa_path.with_suffix(".docx"), [
        ("MASTER SOFTWARE SERVICE AGREEMENT",
         f"Parties: DataVendor LLC (tax id: {COMPANIES['DataVendor']['tax_id']}) and "
         f"GlobalCorp Inc. (tax id: {COMPANIES['GlobalCorp']['tax_id']})\n\n"
         "Effective date: 2026-03-15\nExpiry date: 2027-03-15\nMonthly fee: $35,000 + 20% VAT\n\n"
         "Change of control: 25% ownership change → immediate termination "
         "+ full annual fee ($420,000) as penalty.\n\n"
         "Non-compete: 5 years post-termination.\n\n"
         "Auto-renewal: 3-year extension with 90-day notice.\n\n"
         "Penalty: $2,000 per 1% SLA shortfall."),
    ])
    _render_png_from_pdf(mssa_path, mssa_path.with_suffix(".png"))

    # 3) Amendment — extends the term
    amd_html = f"""
    <h1>AMENDMENT NO. 1 TO SERVICE AGREEMENT</h1>
    <p><b>Original agreement:</b> Master Software Service Agreement dated 2026-03-15
       (DataVendor LLC and GlobalCorp Inc.)</p>
    <p><b>Amendment date:</b> 2026-04-15</p>
    <h2>1. Term extension</h2>
    <p>The expiry date is amended from 2027-03-15 to <b>2028-03-15</b>.</p>
    <h2>2. Monthly fee unchanged</h2>
    <p>The $35,000 + 20% VAT monthly fee remains in effect for the full extended term.</p>
    <h2>3. Other terms</h2>
    <p>The other provisions of the original agreement — including change-of-control,
       non-compete, and auto-renewal clauses — remain unchanged and in full force.</p>
    """
    amd_path = out_dir / "gc-dv-mssa-amd1-2026-0415.pdf"
    _render_html_pdf(amd_path, amd_html)
    _render_docx(amd_path.with_suffix(".docx"), [
        ("AMENDMENT NO. 1 TO SERVICE AGREEMENT",
         "Original agreement: Master Software Service Agreement dated 2026-03-15\n"
         "Amendment date: 2026-04-15\n\n"
         "Term extension: 2027-03-15 → 2028-03-15\n"
         "Monthly fee unchanged: $35,000 + 20% VAT\n\n"
         "All other provisions of the original agreement remain unchanged."),
    ])
    _render_png_from_pdf(amd_path, amd_path.with_suffix(".png"))


def generate_compliance_demo() -> None:
    """Compliance demo: 2 contracts — one with full GDPR Article 28, one missing it."""
    out_dir = THIS_DIR / "demo_packages" / "compliance_demo"

    # 1) Contract A: contains the full GDPR Article 28 clause
    ok_html = f"""
    <h1>DATA PROCESSING AGREEMENT (DPA)</h1>
    <p><b>Parties:</b> MediCare Inc. (tax id: {COMPANIES['MediCare']['tax_id']}) as Controller,
       and CleanLaw LLP (tax id: {COMPANIES['CleanLaw']['tax_id']}) as Processor.</p>
    <p><b>Effective date:</b> 2026-04-01 &nbsp;&nbsp; <b>Expiry date:</b> 2027-04-01</p>
    <h2>1. Scope</h2>
    <p>The Processor processes patient data on behalf of the Controller (health data,
       a special category under GDPR Article 9).</p>
    <h2>2. GDPR Article 28 provisions</h2>
    <p>(a) <b>Subject and duration:</b> The term of this agreement and the duration of the
       service related to processing of patient data.</p>
    <p>(b) <b>Nature and purpose of processing:</b> Patient registry, clinical data storage.</p>
    <p>(c) <b>Type of personal data:</b> Name, ID, address, health records.</p>
    <p>(d) <b>Categories of data subjects:</b> The Controller's patients.</p>
    <p>(e) <b>Controller's rights and duties</b> per GDPR Article 28(3).</p>
    <p>(f) <b>Documented instructions</b> are required for any processing.</p>
    <p>(g) <b>Confidentiality:</b> All Processor staff are under a confidentiality obligation.</p>
    <p>(h) <b>Security measures</b> per GDPR Article 32 (encryption, access control).</p>
    <p>(i) <b>Sub-processor</b> engagement only with prior written consent.</p>
    <p>(j) <b>Assistance</b> with data-subject rights requests.</p>
    <p>(k) <b>Deletion or return</b> of personal data on contract termination.</p>
    <h2>3. Governing law</h2>
    <p>EU data protection law (GDPR) and the laws of the State of New York, USA.</p>
    """
    ok_path = out_dir / "mc-cl-dpa-2026-0401.pdf"
    _render_html_pdf(ok_path, ok_html)
    _render_docx(ok_path.with_suffix(".docx"), [
        ("DATA PROCESSING AGREEMENT (DPA)",
         f"Parties: MediCare Inc. (tax id: {COMPANIES['MediCare']['tax_id']}) as Controller, "
         f"CleanLaw LLP (tax id: {COMPANIES['CleanLaw']['tax_id']}) as Processor\n\n"
         "Effective: 2026-04-01 — Expiry: 2027-04-01\n\n"
         "Full GDPR Article 28 clauses:\n"
         "(a) Subject and duration\n(b) Nature and purpose\n(c) Type of personal data\n"
         "(d) Categories of data subjects\n(e) Controller's rights\n(f) Documented instructions\n"
         "(g) Confidentiality\n(h) Security measures (Art. 32)\n(i) Sub-processor consent\n"
         "(j) Data-subject rights assistance\n(k) Deletion / return of data\n\n"
         "Governing law: GDPR + State of New York, USA."),
    ])
    _render_png_from_pdf(ok_path, ok_path.with_suffix(".png"))

    # 2) Contract B: NO GDPR Article 28 clause despite processing PII
    bad_html = f"""
    <h1>SERVICE AGREEMENT</h1>
    <p><b>Parties:</b> MediCare Inc. (tax id: {COMPANIES['MediCare']['tax_id']}) and
       DataVendor LLC (tax id: {COMPANIES['DataVendor']['tax_id']}).</p>
    <p><b>Effective date:</b> 2026-04-10 &nbsp;&nbsp; <b>Expiry date:</b> 2027-04-10</p>
    <p><b>Monthly fee:</b> $8,000 + 20% VAT</p>
    <h2>1. Scope</h2>
    <p>DataVendor LLC operates a patient registry on behalf of MediCare Inc., including
       patient names, addresses, social security numbers, and clinical data.</p>
    <h2>2. Data processing</h2>
    <p>During the service DataVendor LLC processes personal data (special-category
       health data) on behalf of the controller.</p>
    <h2>3. Liability</h2>
    <p>DataVendor LLC is responsible for the secure handling of the data within its own
       liability framework.</p>
    <h2>4. Governing law</h2>
    <p>State of New York, USA.</p>
    """
    bad_path = out_dir / "mc-dv-msa-2026-0410.pdf"
    _render_html_pdf(bad_path, bad_html)
    _render_docx(bad_path.with_suffix(".docx"), [
        ("SERVICE AGREEMENT",
         f"Parties: MediCare Inc. (tax id: {COMPANIES['MediCare']['tax_id']}) and "
         f"DataVendor LLC (tax id: {COMPANIES['DataVendor']['tax_id']})\n\n"
         "Effective: 2026-04-10 — Expiry: 2027-04-10\nMonthly fee: $8,000 + 20% VAT\n\n"
         "Scope: patient registry (name, SSN, address, clinical data).\n\n"
         "Data security is the Provider's own responsibility within its general liability framework.\n\n"
         "Governing law: State of New York, USA."),
    ])
    _render_png_from_pdf(bad_path, bad_path.with_suffix(".png"))


# ---------------------------------------------------------------------------
# Adversarial — deliberately broken documents
# ---------------------------------------------------------------------------


def generate_adversarial() -> None:
    """4 deliberately broken documents to validate detection logic."""
    out_dir = THIS_DIR / "adversarial"

    # 1) Invoice with math error ($760 difference between net+VAT and gross)
    net_correct = 55_000
    vat_correct = 11_000  # 55_000 × 0.20
    gross_wrong = 67_500  # net+VAT = 66,000 actually, but invoice says 67,500
    inv_html = f"""
    <h1>INVOICE</h1>
    <p><b>Invoice number:</b> ME-2026/0001 &nbsp;&nbsp;
       <b>Issue date:</b> 2026-03-15 &nbsp;&nbsp;
       <b>Fulfillment date:</b> 2026-03-10 &nbsp;&nbsp;
       <b>Payment due:</b> 2026-04-14</p>
    <h2>Issuer</h2>
    <p>Alpha Logistics Inc.<br/>Tax ID: {COMPANIES['AcmeSoft']['tax_id']}<br/>
       Address: 555 Logistics Way, Chicago, IL 60616, USA</p>
    <h2>Customer</h2>
    <p>BuilderInc Inc.<br/>Tax ID: {COMPANIES['BuilderInc']['tax_id']}<br/>
       Address: 1500 Industrial Blvd, Chicago, IL 60616, USA</p>
    <h2>Line items</h2>
    <table>
        <tr><th>Description</th><th>Qty</th><th>Unit price</th><th>Net</th><th>VAT</th></tr>
        <tr><td>Warehouse rental (monthly)</td><td class="right">1 mo</td><td class="right">$15,000.00</td><td class="right">$15,000.00</td><td class="right">20%</td></tr>
        <tr><td>Logistics services</td><td class="right">1 mo</td><td class="right">$28,000.00</td><td class="right">$28,000.00</td><td class="right">20%</td></tr>
        <tr><td>Shipping cost</td><td class="right">1 lot</td><td class="right">$12,000.00</td><td class="right">$12,000.00</td><td class="right">20%</td></tr>
    </table>
    <table>
        <tr><td><b>Total net</b></td><td class="right">{_money(net_correct)}</td></tr>
        <tr><td><b>Total VAT</b></td><td class="right">{_money(vat_correct)}</td></tr>
        <tr class="total"><td><b>Total gross</b></td><td class="right">{_money(gross_wrong)}</td></tr>
    </table>
    """
    inv_path = out_dir / "adv-inv-2026-0001.pdf"
    _render_html_pdf(inv_path, inv_html)
    _render_docx(inv_path.with_suffix(".docx"), [
        ("INVOICE",
         "Invoice number: ME-2026/0001\nIssue: 2026-03-15\nFulfillment: 2026-03-10\nPayment due: 2026-04-14"),
        ("Issuer", f"Alpha Logistics Inc.\nTax ID: {COMPANIES['AcmeSoft']['tax_id']}"),
        ("Customer", f"BuilderInc Inc.\nTax ID: {COMPANIES['BuilderInc']['tax_id']}"),
        ("Line items",
         "Warehouse rental -- 1 mo -- $15,000.00 -- 20%\n"
         "Logistics services -- 1 mo -- $28,000.00 -- 20%\n"
         "Shipping cost -- 1 lot -- $12,000.00 -- 20%"),
        ("Totals",
         f"Net: {_money(net_correct)}\n"
         f"VAT: {_money(vat_correct)}\n"
         f"Gross: {_money(gross_wrong)}"),
    ])
    _render_png_from_pdf(inv_path, inv_path.with_suffix(".png"))

    # 2) Incomplete contract (no termination, no penalty, no expiry date)
    incomplete_html = f"""
    <h1>SERVICE AGREEMENT</h1>
    <p><b>Contract number:</b> HI-2026-001 &nbsp;&nbsp;
       <b>Signing date:</b> 2026-02-15</p>
    <p><b>Parties:</b> Gyros Longrun LLC (tax id: {COMPANIES['AcmeSoft']['tax_id']}) and
       Provider Inc. (tax id: {COMPANIES['DataPharm']['tax_id']}).</p>
    <h2>1. Scope</h2>
    <p>Provision of web development services.</p>
    <h2>2. Compensation</h2>
    <p>$12,000 monthly + 20% VAT.</p>
    """
    inc_path = out_dir / "adv-ctr-2026-001.pdf"
    _render_html_pdf(inc_path, incomplete_html)
    _render_docx(inc_path.with_suffix(".docx"), [
        ("SERVICE AGREEMENT",
         f"Contract number: HI-2026-001\nSigning date: 2026-02-15\n\n"
         "Parties: Gyros Longrun LLC and Provider Inc.\n\n"
         "Scope: Web development.\nMonthly fee: $12,000 + 20% VAT."),
    ])
    _render_png_from_pdf(inc_path, inc_path.with_suffix(".png"))

    # 3) Bilingual contract (HU + EN) with Incoterms CIP
    bilingual_html = f"""
    <h1>SUPPLY AGREEMENT / SZÁLLÍTÁSI SZERZŐDÉS</h1>
    <p><b>Contract number:</b> ML-2026-001 &nbsp;&nbsp;
       <b>Signed:</b> 2026-03-10</p>
    <p><b>Parties / Felek:</b> GlobalCorp Ltd. (UK) and Magyar Industrial Park Inc.
       (tax id: {COMPANIES['ConstructLLC']['tax_id']}).</p>
    <p><b>Effective / Hatály:</b> 2026-04-01 -- 2027-03-31</p>
    <p><b>Value / Érték:</b> 450,000 EUR</p>
    <h2>1. Delivery terms / Szállítási feltételek</h2>
    <p>CIP Budapest (Incoterms 2020). The Supplier bears insurance and shipping cost.</p>
    <h2>2. Termination / Felmondás</h2>
    <p>Either party may terminate with 90 days notice. / Bármely fél 90 nappal előre felmondhat.</p>
    <h2>3. Governing law / Irányadó jog</h2>
    <p>English law. / Angol jog.</p>
    """
    bil_path = out_dir / "adv-ctr-2026-002.pdf"
    _render_html_pdf(bil_path, bilingual_html)
    _render_docx(bil_path.with_suffix(".docx"), [
        ("SUPPLY AGREEMENT / SZÁLLÍTÁSI SZERZŐDÉS",
         f"Contract number: ML-2026-001\nSigned: 2026-03-10\n\n"
         f"Parties: GlobalCorp Ltd. and Magyar Industrial Park Inc.\n\n"
         "Effective: 2026-04-01 -- 2027-03-31\nValue: 450,000 EUR\n\n"
         "Delivery: CIP Budapest (Incoterms 2020).\n"
         "Termination: 90 days notice.\nGoverning law: English law."),
    ])
    _render_png_from_pdf(bil_path, bil_path.with_suffix(".png"))

    # 4) Contract with date-logic errors (signing after start, expiry before start)
    illogical_html = f"""
    <h1>WORK AGREEMENT</h1>
    <p><b>Contract number:</b> ED-2026-001 &nbsp;&nbsp;
       <b>Signing date:</b> 2026-02-15</p>
    <p><b>Parties:</b> Spring Autoservice Inc. (tax id: {COMPANIES['TechSupply']['tax_id']}) and
       Customer Corp. (tax id: {COMPANIES['AcmeBuy']['tax_id']}).</p>
    <p><b>Effective date:</b> 2026-01-01 (back-dated)</p>
    <p><b>Expiry date:</b> 2025-12-15</p>
    <p><b>Fulfillment date:</b> 2025-12-15</p>
    <p><b>Payment due:</b> 2026-02-01 (before signing!)</p>
    <h2>1. Scope</h2>
    <p>Vehicle maintenance services.</p>
    <h2>2. Fee</h2>
    <p>$8,000 + 20% VAT.</p>
    """
    ill_path = out_dir / "adv-ctr-2026-003.pdf"
    _render_html_pdf(ill_path, illogical_html)
    _render_docx(ill_path.with_suffix(".docx"), [
        ("WORK AGREEMENT",
         "Contract number: ED-2026-001\nSigning date: 2026-02-15\n\n"
         "Parties: Spring Autoservice Inc. and Customer Corp.\n\n"
         "Effective date: 2026-01-01 (back-dated)\nExpiry date: 2025-12-15 (BEFORE start!)\n"
         "Fulfillment: 2025-12-15\nPayment due: 2026-02-01 (BEFORE signing!)\n\n"
         "Scope: Vehicle maintenance.\nFee: $8,000 + 20% VAT."),
    ])
    _render_png_from_pdf(ill_path, ill_path.with_suffix(".png"))


# ---------------------------------------------------------------------------
# Financial reports (multilingual demo)
# ---------------------------------------------------------------------------


def generate_financial_reports() -> None:
    """1 EN income statement (US-GAAP) + 1 EN cash flow (IFRS, multilingual demo)."""
    out_dir = THIS_DIR / "financial_reports"

    # 1) Income statement (US-GAAP)
    is_html = f"""
    <h1>INCOME STATEMENT</h1>
    <p><b>Company:</b> FutureTech Inc. (tax id: {COMPANIES['AcmeSoft']['tax_id']})</p>
    <p><b>Period:</b> 2025-01-01 to 2025-12-31 (audited)</p>
    <p><b>Standard:</b> US-GAAP</p>
    <p><b>Currency:</b> USD (thousands)</p>
    <h2>Revenue</h2>
    <table>
        <tr><th>Item</th><th>2025 (kUSD)</th><th>2024 (kUSD)</th></tr>
        <tr><td>Net sales revenue</td><td class="right">485,000</td><td class="right">412,000</td></tr>
        <tr><td>Other income</td><td class="right">12,500</td><td class="right">8,700</td></tr>
        <tr class="total"><td><b>Total revenue</b></td><td class="right">497,500</td><td class="right">420,700</td></tr>
    </table>
    <h2>Costs</h2>
    <table>
        <tr><td>Cost of goods sold</td><td class="right">187,200</td><td class="right">165,100</td></tr>
        <tr><td>Personnel costs</td><td class="right">154,800</td><td class="right">132,400</td></tr>
        <tr><td>Depreciation</td><td class="right">28,600</td><td class="right">31,200</td></tr>
        <tr class="total"><td><b>Operating costs total</b></td><td class="right">370,600</td><td class="right">328,700</td></tr>
    </table>
    <p><b>Operating income (EBIT):</b> 126,900 kUSD (2024: 92,000, +37.9%)</p>
    <p><b>Pretax income:</b> 122,400 kUSD</p>
    """
    is_path = out_dir / "fin-en-is-2025.pdf"
    _render_html_pdf(is_path, is_html)
    _render_docx(is_path.with_suffix(".docx"), [
        ("INCOME STATEMENT",
         "Company: FutureTech Inc.\nPeriod: 2025-01-01 to 2025-12-31\nStandard: US-GAAP\n\n"
         "Total revenue 2025: 497,500 kUSD (2024: 420,700 kUSD, +18.3%)\n"
         "Operating costs 2025: 370,600 kUSD (2024: 328,700 kUSD)\n"
         "EBIT 2025: 126,900 kUSD (2024: 92,000, +37.9%)\n"
         "Pretax income: 122,400 kUSD"),
    ])
    _render_png_from_pdf(is_path, is_path.with_suffix(".png"))

    # 2) Cash flow (IFRS, Alpine Biotech AG)
    cf_html = """
    <h1>CASH FLOW STATEMENT</h1>
    <p><b>Company:</b> Alpine Biotech AG (Switzerland)</p>
    <p><b>Period:</b> 2025-01-01 to 2025-12-31 (audited)</p>
    <p><b>Standard:</b> IFRS (International Financial Reporting Standards)</p>
    <p><b>Currency:</b> CHF (thousands)</p>
    <h2>Operating activities</h2>
    <table>
        <tr><th>Item</th><th>2025 (kCHF)</th></tr>
        <tr><td>Net income</td><td class="right">42,800</td></tr>
        <tr><td>Depreciation &amp; amortization</td><td class="right">18,200</td></tr>
        <tr><td>Working capital changes</td><td class="right">-3,400</td></tr>
        <tr class="total"><td><b>Cash from operating</b></td><td class="right">57,600</td></tr>
    </table>
    <h2>Investing activities</h2>
    <table>
        <tr><td>Production line CapEx</td><td class="right">-67,400</td></tr>
        <tr><td>R&amp;D investments</td><td class="right">-12,100</td></tr>
        <tr class="total"><td><b>Cash from investing</b></td><td class="right">-79,500</td></tr>
    </table>
    <h2>Financing activities</h2>
    <table>
        <tr><td>Bond issuance (5y, 4.2%)</td><td class="right">35,000</td></tr>
        <tr><td>Dividend paid</td><td class="right">-8,200</td></tr>
        <tr class="total"><td><b>Cash from financing</b></td><td class="right">26,800</td></tr>
    </table>
    <p><b>Net change in cash:</b> 4,900 kCHF</p>
    """
    cf_path = out_dir / "fin-en-cf-2025.pdf"
    _render_html_pdf(cf_path, cf_html)
    _render_docx(cf_path.with_suffix(".docx"), [
        ("CASH FLOW STATEMENT",
         "Company: Alpine Biotech AG\nPeriod: 2025-01-01 to 2025-12-31\nStandard: IFRS\nCurrency: CHF\n\n"
         "Operating: +57,600 kCHF\n"
         "Investing: -79,500 kCHF (Production CapEx -67,400, R&D -12,100)\n"
         "Financing: +26,800 kCHF (Bond 35,000, Dividend -8,200)\n"
         "Net change in cash: +4,900 kCHF"),
    ])
    _render_png_from_pdf(cf_path, cf_path.with_suffix(".png"))


# ---------------------------------------------------------------------------
# Cleanup
# ---------------------------------------------------------------------------


def _cleanup_test_data_dirs() -> None:
    """Clear stale generated files before regeneration."""
    target_dirs = [
        THIS_DIR / "invoices",
        THIS_DIR / "contracts",
        THIS_DIR / "multi_doc",
        THIS_DIR / "adversarial",
        THIS_DIR / "financial_reports",
        THIS_DIR / "demo_packages" / "audit_demo",
        THIS_DIR / "demo_packages" / "dd_demo",
        THIS_DIR / "demo_packages" / "compliance_demo",
    ]
    for d in target_dirs:
        d.mkdir(parents=True, exist_ok=True)
    deleted = 0
    for d in target_dirs:
        if not d.exists():
            continue
        for ext in (".pdf", ".docx", ".png"):
            for f in d.glob(f"*{ext}"):
                f.unlink()
                deleted += 1
    if deleted > 0:
        print(f"  Cleanup: {deleted} stale files removed")


def main() -> None:
    print("Generating sample data...")
    _cleanup_test_data_dirs()
    generate_invoices()
    print("  3 EN invoices (audit pattern: March 50% pricier)")
    generate_intra_eu_invoice()
    print("  1 EN intra-EU invoice (0% VAT, reverse charge)")
    generate_de_rechnung()
    print("  1 DE Rechnung (19% MwSt, multilingual demo)")
    generate_contracts()
    print("  4 contracts (NDA + MSSA + IT framework + DE→HU lease)")
    generate_multi_doc_triplet()
    print("  3 multi_doc (PO + delivery_note + invoice with quantity discrepancy)")
    generate_audit_demo()
    print("  Audit demo package (3 invoices, 50% price increase)")
    generate_dd_demo()
    print("  DD demo package (NDA + MSSA + amendment)")
    generate_compliance_demo()
    print("  Compliance demo package (2 contracts, GDPR asymmetry)")
    generate_adversarial()
    print("  4 adversarial documents (math error, incomplete, bilingual, date errors)")
    generate_financial_reports()
    print("  2 financial reports (US-GAAP IS + IFRS CF)")

    pdf_count = sum(1 for _ in THIS_DIR.rglob("*.pdf"))
    docx_count = sum(1 for _ in THIS_DIR.rglob("*.docx"))
    png_count = sum(1 for _ in THIS_DIR.rglob("*.png"))
    print(f"\nTotal: {pdf_count} PDF, {docx_count} DOCX, {png_count} PNG")


if __name__ == "__main__":
    main()
